import os
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path

from ..server import mcp, client
from ._annotations import LOCAL_READ_ONLY, READ_ONLY

ALLOWED_EXTENSIONS = {".html", ".htm", ".md", ".txt", ".csv", ".json"}
OPTIONAL_EXTENSIONS = {".pdf", ".docx"}
ALL_EXTENSIONS = ALLOWED_EXTENSIONS | OPTIONAL_EXTENSIONS

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_LENGTH = 50_000
MAX_SCAN_FILES = 200

SKIP_HTML_TAGS = {"script", "style", "nav", "footer"}


# ---------------------------------------------------------------------------
# HTML text extraction via stdlib
# ---------------------------------------------------------------------------

class _HTMLTextExtractor(HTMLParser):
    """Minimal HTML-to-text extractor using stdlib html.parser."""

    def __init__(self):
        super().__init__()
        self._pieces: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs):
        if tag.lower() in SKIP_HTML_TAGS:
            self._skip_depth += 1

    def handle_endtag(self, tag: str):
        if tag.lower() in SKIP_HTML_TAGS:
            self._skip_depth = max(0, self._skip_depth - 1)

    def handle_data(self, data: str):
        if self._skip_depth == 0:
            text = data.strip()
            if text:
                self._pieces.append(text)

    def get_text(self) -> str:
        return "\n".join(self._pieces)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _get_content_dir() -> Path | None:
    """Return AEKO_CONTENT_DIR as a resolved Path, or None if not set."""
    raw = os.environ.get("AEKO_CONTENT_DIR", "")
    if not raw:
        return None
    p = Path(raw).resolve()
    if not p.is_dir():
        raise RuntimeError(f"AEKO_CONTENT_DIR path does not exist or is not a directory: {p}")
    return p


def _validate_content_path(file_path: str) -> Path:
    """Validate a content file path: extension whitelist, size cap, sandbox check."""
    p = Path(file_path)
    content_dir = _get_content_dir()

    # Resolve relative paths against content dir
    if not p.is_absolute():
        if content_dir is None:
            raise ValueError(
                "Relative paths require AEKO_CONTENT_DIR to be set. "
                "Provide an absolute path or set the environment variable."
            )
        p = content_dir / p

    p = p.resolve()

    # Sandbox check: if AEKO_CONTENT_DIR is set, path must be within it
    if content_dir is not None:
        if not str(p).startswith(str(content_dir)):
            raise ValueError(
                f"Path must be within AEKO_CONTENT_DIR ({content_dir}). Got: {p}"
            )

    # Extension check
    if p.suffix.lower() not in ALL_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {p.suffix}. "
            f"Allowed: {', '.join(sorted(ALL_EXTENSIONS))}"
        )

    # Existence check
    if not p.is_file():
        raise FileNotFoundError(f"File not found: {p}")

    # Size check
    size = p.stat().st_size
    if size > MAX_FILE_SIZE:
        raise ValueError(
            f"File too large: {size / 1024 / 1024:.1f} MB (limit: {MAX_FILE_SIZE / 1024 / 1024:.0f} MB)"
        )

    return p


def _extract_html(path: Path) -> tuple[str, str]:
    """Extract text and raw HTML from an HTML file."""
    raw_html = path.read_text(encoding="utf-8", errors="replace")
    extractor = _HTMLTextExtractor()
    extractor.feed(raw_html)
    return extractor.get_text(), raw_html


def _extract_text(path: Path) -> tuple[str, dict]:
    """Read a content file and return (text, metadata).

    Returns extracted text and a metadata dict with format-specific info.
    """
    ext = path.suffix.lower()
    meta: dict = {}

    if ext in (".html", ".htm"):
        text, raw_html = _extract_html(path)
        meta["raw_html_length"] = len(raw_html)
    elif ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise RuntimeError(
                "PyPDF2 is required to read PDF files. "
                "Install it with: pip install 'aeko-mcp[pdf]'"
            )
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)
        text = "\n\n".join(pages)
        meta["page_count"] = len(reader.pages)
    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx is required to read DOCX files. "
                "Install it with: pip install 'aeko-mcp[docx]'"
            )
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        meta["paragraph_count"] = len(doc.paragraphs)
    else:
        # .md, .txt, .csv, .json — read as plain text
        text = path.read_text(encoding="utf-8", errors="replace")

    return text, meta


def _file_metadata(path: Path) -> dict:
    """Return basic metadata about a content file."""
    stat = path.stat()
    text, _ = _extract_text(path)
    word_count = len(text.split()) if text else 0

    return {
        "name": path.name,
        "path": str(path),
        "extension": path.suffix.lower(),
        "size_bytes": stat.st_size,
        "size_display": (
            f"{stat.st_size / 1024:.1f} KB" if stat.st_size < 1024 * 1024
            else f"{stat.st_size / 1024 / 1024:.1f} MB"
        ),
        "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).strftime("%Y-%m-%d %H:%M"),
        "word_count": word_count,
    }


# ---------------------------------------------------------------------------
# Tools
# ---------------------------------------------------------------------------

@mcp.tool(annotations=LOCAL_READ_ONLY)
def aeko_scan_content_directory(
    directory: str = "",
    extensions: str = "",
    recursive: bool = True,
) -> str:
    """Scan a local directory for content files (HTML, Markdown, text, CSV, JSON).

    Lists content files with metadata. Use this to discover files before
    reading or auditing them. Defaults to AEKO_CONTENT_DIR if set.

    Args:
        directory: Path to scan. Defaults to AEKO_CONTENT_DIR env var.
        extensions: Comma-separated list of extensions to filter (e.g. ".html,.md"). Defaults to all supported types.
        recursive: Whether to scan subdirectories. Defaults to True.

    Returns:
        Markdown table of content files with type, size, and modification date.
    """
    if directory:
        base = Path(directory).resolve()
        if not base.is_dir():
            return f"Error: directory not found: {directory}"
        # Sandbox check
        content_dir = _get_content_dir()
        if content_dir is not None and not str(base).startswith(str(content_dir)):
            return f"Error: path must be within AEKO_CONTENT_DIR ({content_dir})"
    else:
        content_dir = _get_content_dir()
        if content_dir is None:
            return (
                "Error: no directory specified and AEKO_CONTENT_DIR is not set. "
                "Provide a directory path or set the AEKO_CONTENT_DIR environment variable."
            )
        base = content_dir

    # Parse extensions filter
    if extensions:
        ext_filter = {e.strip().lower() if e.strip().startswith(".") else f".{e.strip().lower()}" for e in extensions.split(",")}
        invalid = ext_filter - ALL_EXTENSIONS
        if invalid:
            return f"Error: unsupported extension(s): {', '.join(sorted(invalid))}. Allowed: {', '.join(sorted(ALL_EXTENSIONS))}"
    else:
        ext_filter = ALL_EXTENSIONS

    # Collect files
    files: list[Path] = []
    for ext in ext_filter:
        pattern = f"*{ext}"
        if recursive:
            files.extend(base.rglob(pattern))
        else:
            files.extend(base.glob(pattern))

    if not files:
        return f"No content files found in {base}"

    # Sort by modification time (newest first), cap at MAX_SCAN_FILES
    files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    total = len(files)
    files = files[:MAX_SCAN_FILES]

    lines = [
        f"# Content Files in {base}",
        f"Found {total} file(s){f' (showing first {MAX_SCAN_FILES})' if total > MAX_SCAN_FILES else ''}.",
        "",
        "| File | Type | Size | Modified | Words |",
        "|------|------|------|----------|-------|",
    ]

    for f in files:
        stat = f.stat()
        size = f"{stat.st_size / 1024:.1f} KB" if stat.st_size < 1024 * 1024 else f"{stat.st_size / 1024 / 1024:.1f} MB"
        modified = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).strftime("%Y-%m-%d %H:%M")
        rel = f.relative_to(base)
        ext = f.suffix.lower()

        # Quick word count (read text only for small files to keep scan fast)
        words = "—"
        if stat.st_size < 512 * 1024:  # 512 KB threshold for word count
            try:
                text, _ = _extract_text(f)
                words = str(len(text.split())) if text else "0"
            except Exception:
                words = "?"

        lines.append(f"| `{rel}` | {ext} | {size} | {modified} | {words} |")

    return "\n".join(lines)


@mcp.tool(annotations=LOCAL_READ_ONLY)
def aeko_read_content_file(file_path: str, include_raw_html: bool = False) -> str:
    """Read a local content file and extract its text content.

    Supports HTML (strips tags), Markdown, text, CSV, JSON.
    PDF and DOCX supported if optional dependencies are installed.

    Args:
        file_path: Path to the content file (absolute, or relative to AEKO_CONTENT_DIR).
        include_raw_html: For HTML files, also include the raw HTML source. Default False.

    Returns:
        Extracted text content with file metadata.
    """
    path = _validate_content_path(file_path)
    ext = path.suffix.lower()

    # Extract text
    text, extra_meta = _extract_text(path)

    # Truncate if needed
    truncated = False
    if len(text) > MAX_TEXT_LENGTH:
        text = text[:MAX_TEXT_LENGTH]
        truncated = True

    # Build output
    stat = path.stat()
    word_count = len(text.split()) if text else 0
    modified = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        f"# {path.name}",
        "",
        f"- **Path**: `{path}`",
        f"- **Type**: {ext}",
        f"- **Size**: {stat.st_size / 1024:.1f} KB",
        f"- **Modified**: {modified}",
        f"- **Words**: {word_count:,}",
    ]

    for key, val in extra_meta.items():
        lines.append(f"- **{key.replace('_', ' ').title()}**: {val}")

    if truncated:
        lines.append(f"- **Note**: Content truncated at {MAX_TEXT_LENGTH:,} characters")

    lines.append("")
    lines.append("## Content")
    lines.append("")
    lines.append(text)

    # Optionally include raw HTML
    if include_raw_html and ext in (".html", ".htm"):
        raw_html = path.read_text(encoding="utf-8", errors="replace")
        if len(raw_html) > MAX_TEXT_LENGTH:
            raw_html = raw_html[:MAX_TEXT_LENGTH] + "\n\n[... truncated]"
        lines.append("")
        lines.append("## Raw HTML")
        lines.append("```html")
        lines.append(raw_html)
        lines.append("```")

    return "\n".join(lines)


@mcp.tool(annotations=READ_ONLY)
def aeko_audit_content_file(file_path: str, language: str = "") -> str:
    """Read a local content file and score it for AI citability.

    Combines local file reading with AEKO's citability scoring API.
    Returns file metadata, citability score breakdown (5 dimensions),
    and improvement suggestions.

    Args:
        file_path: Path to the content file (absolute, or relative to AEKO_CONTENT_DIR).
        language: Optional ISO language code (e.g. 'en', 'ko') for language-specific analysis.

    Returns:
        Combined file metadata and citability analysis with scores and recommendations.
    """
    path = _validate_content_path(file_path)

    # Extract text
    text, _ = _extract_text(path)
    if not text.strip():
        return f"Error: file is empty or contains no extractable text: {path.name}"

    # Truncate for API call
    api_text = text[:MAX_TEXT_LENGTH]

    # File metadata
    stat = path.stat()
    word_count = len(text.split())
    modified = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Call citability API
    payload: dict = {"text": api_text}
    if language:
        payload["language"] = language

    try:
        data = client.post("/api/citability/score", json=payload)
    except Exception as e:
        return (
            f"# {path.name} — Citability Audit\n\n"
            f"File read successfully ({word_count:,} words) but citability scoring failed:\n"
            f"{e}\n\n"
            "Make sure the AEKO backend is running and accessible."
        )

    overall = data.get("overall", 0)
    dims = data.get("dimensions", {})
    improvements = data.get("top_improvements", [])

    # Grade
    if overall >= 90:
        grade = "A"
    elif overall >= 75:
        grade = "B"
    elif overall >= 60:
        grade = "C"
    elif overall >= 40:
        grade = "D"
    else:
        grade = "F"

    # Build report
    lines = [
        f"# Citability Audit: {path.name}",
        "",
        "## File Info",
        f"- **Path**: `{path}`",
        f"- **Type**: {path.suffix.lower()}",
        f"- **Size**: {stat.st_size / 1024:.1f} KB",
        f"- **Words**: {word_count:,}",
        f"- **Modified**: {modified}",
        "",
        f"## Citability Score: {overall}/100 (Grade: {grade})",
        "",
        "| Dimension | Score | Weight |",
        "|-----------|-------|--------|",
    ]

    dim_labels = {
        "answer_block_quality": "Answer Block Quality",
        "self_containment": "Self-Containment",
        "structural_readability": "Structural Readability",
        "statistical_density": "Statistical Density",
        "uniqueness_signals": "Uniqueness Signals",
    }

    for key, label in dim_labels.items():
        d = dims.get(key, {})
        lines.append(f"| {label} | {d.get('score', 0)}/100 | {int(d.get('weight', 0) * 100)}% |")

    if improvements:
        lines.append("")
        lines.append("## Top Improvements")
        for imp in improvements:
            lines.append(f"- {imp}")

    return "\n".join(lines)
